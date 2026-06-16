"""Exam Attempts Export Service

Exports all submitted (DONE) user test attempts with scores (Bảng điểm)
for a given exam_template_id or exam_instance_id as a Word (.docx) document.
"""

from __future__ import annotations

import io
import logging
from datetime import datetime
from uuid import UUID

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from src.entities.exam_template import ExamTemplate
from src.entities.user_test_attempt import UserTestAttempt
from src.repos.exam_instance_repo import ExamInstanceRepository
from src.repos.user_test_attempt_repo import UserTestAttemptRepository

logger = logging.getLogger(__name__)

_HEADERS = ["STT", "Họ và tên", "Email", "Thời gian bắt đầu", "Thời gian nộp", "Điểm"]
_COL_WIDTHS = [Cm(1.2), Cm(4.5), Cm(5.0), Cm(4.0), Cm(4.0), Cm(2.0)]
_CENTER_COLS = {0, 3, 4, 5}
_HEADER_BG = "BDD7EE"
_ROW_ALT_BG = "F2F2F2"


class ExamAttemptsExportService:
    """Builds exam attempt score sheets (Bảng điểm) as Word documents."""

    def __init__(self):
        self._attempt_repo = UserTestAttemptRepository()
        self._instance_repo = ExamInstanceRepository()

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def export_by_template(self, exam_template_id: UUID) -> tuple[bytes, str]:
        """Export submitted attempts for all instances of a template."""
        attempts = self._attempt_repo.list_submitted_by_template(exam_template_id)

        try:
            template = ExamTemplate.get_by_id(str(exam_template_id))
            title = f"Bảng Điểm - {template.name}"
            subtitle = f"Môn: {template.subject}"
        except Exception:
            title = "Bảng Điểm"
            subtitle = f"Mã đề: {exam_template_id}"

        doc_bytes = self._build_docx(title=title, subtitle=subtitle, attempts=attempts)
        filename = f"bang_diem_template_{exam_template_id}.docx"
        return doc_bytes, filename

    def export_by_instance(self, exam_instance_id: UUID) -> tuple[bytes, str]:
        """Export submitted attempts for a specific exam instance."""
        attempts = self._attempt_repo.list_submitted_by_instance(exam_instance_id)

        exam_code = str(exam_instance_id)
        try:
            instance = self._instance_repo.get_by_id(exam_instance_id)
            if instance:
                exam_code = instance.exam_test_code or exam_code
        except Exception:
            pass

        title = f"Bảng Điểm - {exam_code}"
        subtitle = f"Mã đề thi: {exam_code}"

        doc_bytes = self._build_docx(title=title, subtitle=subtitle, attempts=attempts)
        filename = f"bang_diem_{exam_code}.docx"
        return doc_bytes, filename

    # ------------------------------------------------------------------
    # Document builder
    # ------------------------------------------------------------------

    def _build_docx(
        self,
        title: str,
        subtitle: str,
        attempts: list[UserTestAttempt],
    ) -> bytes:
        doc = Document()

        section = doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        self._add_title_block(doc, title, subtitle, len(attempts))
        self._add_table(doc, attempts)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _add_title_block(
        self, doc: Document, title: str, subtitle: str, total: int
    ) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(16)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(12)

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(
            f"Xuất ngày: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        )
        r3.font.size = Pt(10)
        r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph()

        p4 = doc.add_paragraph()
        r4 = p4.add_run(f"Tổng số thí sinh đã nộp bài: {total}")
        r4.bold = True
        r4.font.size = Pt(11)

        doc.add_paragraph()

    def _add_table(self, doc: Document, attempts: list[UserTestAttempt]) -> None:
        table = doc.add_table(rows=1, cols=len(_HEADERS))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        hdr = table.rows[0]
        for i, (header, width) in enumerate(zip(_HEADERS, _COL_WIDTHS)):
            cell = hdr.cells[i]
            cell.width = width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(header)
            r.bold = True
            r.font.size = Pt(11)
            _set_cell_bg(cell, _HEADER_BG)

        # Data rows
        for idx, attempt in enumerate(attempts, 1):
            row = table.add_row()
            values = self._attempt_to_row(idx, attempt)

            for col, value in enumerate(values):
                cell = row.cells[col]
                cell.width = _COL_WIDTHS[col]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if col in _CENTER_COLS
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                r = p.add_run(value)
                r.font.size = Pt(10)

                if idx % 2 == 0:
                    _set_cell_bg(cell, _ROW_ALT_BG)

    @staticmethod
    def _attempt_to_row(idx: int, attempt: UserTestAttempt) -> list[str]:
        try:
            user_name = attempt.user.name
            user_email = attempt.user.email
        except Exception:
            user_name = str(getattr(attempt, "user_id", ""))
            user_email = ""

        started_at = (
            attempt.started_at.strftime("%d/%m/%Y %H:%M") if attempt.started_at else ""
        )
        submitted_at = (
            attempt.submitted_at.strftime("%d/%m/%Y %H:%M")
            if attempt.submitted_at
            else ""
        )
        score = f"{float(attempt.score):.2f}" if attempt.score is not None else "—"

        return [str(idx), user_name, user_email, started_at, submitted_at, score]


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)
